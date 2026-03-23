from pathlib import Path
import numpy as np
import pandas as pd
from scipy import stats
import statsmodels.formula.api as smf
from docx import Document
import matplotlib.pyplot as plt

ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
PLOTS = ROOT / "plots"
REPORTS.mkdir(exist_ok=True)
PLOTS.mkdir(exist_ok=True)

ELECTION_YEARS = {2011, 2015, 2020, 2025}


def load_data() -> pd.DataFrame:
    df = pd.read_csv(ROOT / "COEBiddingResultsPrices.csv")
    df["month"] = pd.to_datetime(df["month"])
    for col in ["premium", "quota", "bids_received"]:
        df[col] = pd.to_numeric(df[col], errors="coerce")
    df = df[df["vehicle_class"].isin(["Category A", "Category B", "Category E"])].copy()
    df = df.dropna(subset=["premium", "quota", "bids_received"])
    df["year"] = df["month"].dt.year
    df["election_year"] = df["year"].isin(ELECTION_YEARS).astype(int)
    df["log_premium"] = np.log(df["premium"])
    df["log_quota"] = np.log(df["quota"])
    df["log_bids"] = np.log(df["bids_received"])
    df["time_index"] = (df["month"] - df["month"].min()).dt.days / 30.44
    return df


def run_tests(df: pd.DataFrame) -> dict:
    election = df[df["election_year"] == 1]["log_premium"]
    non_election = df[df["election_year"] == 0]["log_premium"]

    t_stat, p_val = stats.ttest_ind(election, non_election, equal_var=False)

    model_simple = smf.ols(
        "log_premium ~ election_year + C(vehicle_class)", data=df
    ).fit(cov_type="HC3")

    model_controlled = smf.ols(
        "log_premium ~ election_year + log_quota + log_bids + C(vehicle_class)", data=df
    ).fit(cov_type="HC3")

    model_with_trend = smf.ols(
        "log_premium ~ election_year + log_quota + log_bids + time_index + C(vehicle_class)", data=df
    ).fit(cov_type="HC3")

    return {
        "n_total": len(df),
        "n_election": len(election),
        "n_non_election": len(non_election),
        "mean_log_election": float(election.mean()),
        "mean_log_non_election": float(non_election.mean()),
        "mean_price_election": float(np.exp(election.mean())),
        "mean_price_non_election": float(np.exp(non_election.mean())),
        "ttest_stat": float(t_stat),
        "ttest_p": float(p_val),
        "simple_coef": float(model_simple.params.get("election_year", np.nan)),
        "simple_p": float(model_simple.pvalues.get("election_year", np.nan)),
        "controlled_coef": float(model_controlled.params.get("election_year", np.nan)),
        "controlled_p": float(model_controlled.pvalues.get("election_year", np.nan)),
        "trend_coef": float(model_with_trend.params.get("election_year", np.nan)),
        "trend_p": float(model_with_trend.pvalues.get("election_year", np.nan)),
        "r2_simple": float(model_simple.rsquared),
        "r2_controlled": float(model_controlled.rsquared),
        "r2_trend": float(model_with_trend.rsquared),
    }


def save_outputs(df: pd.DataFrame, res: dict) -> None:
    out = pd.DataFrame([res])
    out.to_csv(REPORTS / "rq4_election_test_results.csv", index=False)

    # By-class means for reporting
    class_summary = (
        df.groupby(["vehicle_class", "election_year"])["premium"]
        .mean()
        .reset_index()
        .pivot(index="vehicle_class", columns="election_year", values="premium")
        .rename(columns={0: "non_election_mean", 1: "election_mean"})
        .reset_index()
    )
    class_summary.to_csv(REPORTS / "rq4_election_by_class_means.csv", index=False)

    # Plot
    plt.figure(figsize=(8, 5))
    labels = ["Non-election years", "Election years"]
    values = [res["mean_price_non_election"], res["mean_price_election"]]
    bars = plt.bar(labels, values, color=["#4c78a8", "#f58518"])
    plt.ylabel("Mean COE Premium (SGD)")
    plt.title("RQ4: Mean COE Premium in Election vs Non-election Years\n(Categories A/B/E pooled)")
    for b, v in zip(bars, values):
        plt.text(b.get_x() + b.get_width()/2, v, f"${v:,.0f}", ha="center", va="bottom", fontsize=9)
    plt.tight_layout()
    plt.savefig(PLOTS / "rq4_election_vs_nonelection.png", dpi=150)
    plt.close()

    # Text summary
    summary = [
        "RQ4 Election Effect Test Summary",
        f"N total={res['n_total']}, election={res['n_election']}, non-election={res['n_non_election']}",
        f"Mean premium election years=${res['mean_price_election']:,.0f}",
        f"Mean premium non-election years=${res['mean_price_non_election']:,.0f}",
        f"Welch t-test p-value={res['ttest_p']:.6g}",
        "",
        "Regression election-year coefficient (log premium scale):",
        f"- Simple model (class controls): coef={res['simple_coef']:.4f}, p={res['simple_p']:.6g}",
        f"- Controlled model (+quota +bids): coef={res['controlled_coef']:.4f}, p={res['controlled_p']:.6g}",
        f"- Trend model (+time trend): coef={res['trend_coef']:.4f}, p={res['trend_p']:.6g}",
        "",
        "Interpretation rule:",
        "If p < 0.05 and coef<0, evidence supports election-year price drop.",
        "If p >= 0.05 or coef>=0, no evidence of election-year drop.",
    ]
    (REPORTS / "rq4_election_summary.txt").write_text("\n".join(summary), encoding="utf-8")

    # Short docx report
    doc = Document()
    doc.add_heading("RQ4 Short Report: Do COE Prices Drop in Election Years?", level=1)
    doc.add_paragraph("Data: COE Categories A, B, E pooled; election years = 2011, 2015, 2020, 2025.")

    doc.add_heading("Method", level=2)
    doc.add_paragraph("1) Welch two-sample t-test comparing election-year vs non-election-year mean prices.")
    doc.add_paragraph("2) Regression with election-year dummy and class controls.")
    doc.add_paragraph("3) Controlled regressions adding quota, bids, and a time trend.")

    doc.add_heading("Results", level=2)
    doc.add_paragraph(
        f"Mean premium (election years): ${res['mean_price_election']:,.0f}\n"
        f"Mean premium (non-election years): ${res['mean_price_non_election']:,.0f}\n"
        f"Welch t-test p-value: {res['ttest_p']:.6g}"
    )
    doc.add_paragraph(
        f"Election dummy coefficients (log scale):\n"
        f"Simple: {res['simple_coef']:.4f} (p={res['simple_p']:.6g})\n"
        f"Controlled (+quota+bids): {res['controlled_coef']:.4f} (p={res['controlled_p']:.6g})\n"
        f"Trend model: {res['trend_coef']:.4f} (p={res['trend_p']:.6g})"
    )

    doc.add_heading("Conclusion", level=2)
    if (res["simple_p"] < 0.05 and res["simple_coef"] < 0) or (res["controlled_p"] < 0.05 and res["controlled_coef"] < 0):
        doc.add_paragraph("There is statistical evidence consistent with an election-year price drop.")
    else:
        doc.add_paragraph("There is no robust statistical evidence that COE prices drop in election years.")

    doc.add_paragraph("Note: Election-year dummy is coarse; future work can test pre-election windows (6-12 months) around exact election dates.")
    doc.save(REPORTS / "RQ4_Election_Effect_Short_Report.docx")


def main() -> None:
    df = load_data()
    res = run_tests(df)
    save_outputs(df, res)

    print("Saved:")
    print("- reports/rq4_election_test_results.csv")
    print("- reports/rq4_election_by_class_means.csv")
    print("- reports/rq4_election_summary.txt")
    print("- reports/RQ4_Election_Effect_Short_Report.docx")
    print("- plots/rq4_election_vs_nonelection.png")
    print("\nKey results:")
    print(f"Welch t-test p={res['ttest_p']:.6g}")
    print(f"Election coef (simple)={res['simple_coef']:.4f}, p={res['simple_p']:.6g}")
    print(f"Election coef (controlled)={res['controlled_coef']:.4f}, p={res['controlled_p']:.6g}")
    print(f"Election coef (trend)={res['trend_coef']:.4f}, p={res['trend_p']:.6g}")


if __name__ == "__main__":
    main()
